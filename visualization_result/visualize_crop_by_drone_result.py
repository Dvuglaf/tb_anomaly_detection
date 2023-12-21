import os

import cv2

from __crop_by_drone_center_size import videos_drone_pos

from docx import Document
from docx import shared
from docx.shared import Pt

from matplotlib import pyplot as plt


def get_bbox_of_cropped(frame, c, s):
    s[0] = (s[0] + s[1]) // 2
    s[1] = s[0]
    _x1, _y1, _x2, _y2 = int(c[0] - s[0]), int(c[1] - s[1]), int(c[0] + s[0]), int(c[1] + s[1])

    y1, y2 = max(0, _y1), min(frame.shape[0], _y2)
    x1, x2 = max(0, _x1), min(frame.shape[1], _x2)

    if _y2 > frame.shape[0]:  # not square
        n = frame.shape[0] - _y1
        x_shift = ((_x2 - _x1) - n) // 2
        x1 = _x1 + x_shift
        x2 = _x2 - x_shift
        y2 = frame.shape[0]

    if _y1 < 0:  # not square
        n = y2 - 0
        x_shift = ((x2 - x1) - n) // 2
        x1 = x1 + x_shift
        x2 = x2 - x_shift
        y2 = n

    if (y2 - y1) > (x2 - x1):
        shift = (y2 - y1) - (x2 - x1)
        if shift < 3:
            y2 -= shift
        else:
            print("Aaaaaaaaaaaaaaaaaaaaaaaa")
    elif (x2 - x1) > (y2 - y1):
        shift = (x2 - x1) - (y2 - y1)
        if shift < 3:
            x2 -= shift
        else:
            print("Aaaaaaaaaaaaaaaaaaaaaaaa")
    return x1, x2, y1, y2


def add_text_to_paragraph(paragraph, text, pt=12):
    wr = paragraph.add_run(text)
    wr.font.size = Pt(pt)


def add_paragraph_to_document(document, text, pt=12):
    wp = document.add_paragraph()
    add_text_to_paragraph(wp, text, pt)


def add_pictures_to_document(document, pictures, text, picture_width_in_inches):
    wp = document.add_paragraph()  # paragraph for picture and its caption
    wp.alignment = 1  # center alignment
    wr = wp.add_run()
    for picture in pictures:
        wr.add_picture(picture, width=shared.Inches(picture_width_in_inches))

    wr.add_text(text)
    wr.italic = True
    add_paragraph_to_document(document, '')
    add_paragraph_to_document(document, '')


document = Document()
document.add_heading('Демонстрация работы обрезки видео по дрону', level=1)
add_paragraph_to_document(document, '')


if __name__ == '__main__':
    data_path = '/Users/dvuglaf/Desktop/cropped_frames/'

    for i, image in enumerate(os.listdir(f"{data_path}/original_frames")):
        if "checkpoint" in image or not ".jpg" in image:
            continue

        print(image)

        original_path = f"{data_path}/original_frames/{image}"
        cropped_path = f"{data_path}/cropped_frames/{image}"
        rectangle_path = f"{data_path}/rectangle_frames/{image}"
        filename = image.replace('.TS', '')[:-4]

        orig_frame = cv2.imread(original_path)

        c, s = videos_drone_pos[filename][1], videos_drone_pos[filename][2]

        _x1, _y1, _x2, _y2 = int(c[0] - s[0]), int(c[1] - s[1]), int(c[0] + s[0]), int(c[1] + s[1])
        x1, x2, y1, y2 = get_bbox_of_cropped(orig_frame, c, s)
        crop_frame = orig_frame[y1:y2, x1:x2]

        cv2.imwrite(cropped_path, crop_frame)
        print(cropped_path)

        print(x1, x2, y1, y2)
        resolution = videos_drone_pos[filename][0]
        print(resolution)
        rectangle_frame = cv2.rectangle(orig_frame,
                                        (_x1, _y1),
                                        (_x2, _y2),
                                        (0, 0, 255),
                                        3)

        rectangle_frame = cv2.rectangle(rectangle_frame,
                                        (x1, y1),
                                        (x2, y2),
                                        (0, 255, 0),
                                        3)
        cv2.imwrite(rectangle_path, rectangle_frame)

        orig_shape = cv2.imread(original_path).shape
        crop_shape = cv2.imread(cropped_path).shape

        add_pictures_to_document(
            document,
            pictures=[rectangle_path, cropped_path],
            text=f"{orig_shape}----\"{image.split('/')[-1]}\"---->{crop_shape}",
            picture_width_in_inches=3.
        )

    document.save("test_rectangle_example.docx")


exit(1)
